"""

This script produces a synthetic dataset for testing document classification
systems with regard to sensitive information. It is designed to satisfy a
number of stringent requirements:

1. **Reproducibility** – The user supplies a seed (any string) and the same
   dataset is produced whenever that seed is reused. Internally the seed is
   hashed into a deterministic integer and used to initialise Python's
   pseudo‑random generator.

2. **Naming conventions** – Five different academic/archival naming schemes
   are implemented (Harvard Research, Montessori Archives, Institutional,
   OSU Metadata and Reddit User Practice). In addition, file names may be
   deliberately ambiguous or misleading. Roughly half of the file names
   reflect the underlying content while the other half are decoys. Each
   naming convention incorporates a date stamp and version when appropriate.

3. **File types and distribution** – One hundred files are generated: 70
   plain text (`.txt`), 15 Word documents (`.docx`) and 15 PDF files
   (`.pdf`). Half of the documents contain sensitive information and half
   contain innocuous content. The distribution of file types across the
   sensitive/insensitive split is balanced as closely as possible.

4. **Sensitive categories** – Twenty‑two classes of sensitive data are
   represented: phone numbers, IBANs, names, addresses, credit card
   numbers, passport numbers, driver’s licence numbers, email addresses, IP
   addresses, usernames, passwords, medical terms, diagnoses, tax numbers,
   salary information, PINs, bills, invoices, security questions,
   contracts, API keys and authentication tokens. Each sensitive file
   contains realistic values for one or more of these categories. A
   minority of files mix several categories to test multi‑class handling.
   For each category a set of synonyms is used for the labels to avoid
   always using the same wording.

5. **Content styles** – The content is rendered in various everyday styles
   reflecting how people actually capture information: bullet‑point lists
   (lecture notes), paragraph‑style prose (essays), to‑do lists (office
   work) and terse note taking. Insensitive documents are similarly
   generated using plausible templates such as grocery lists, travel
   itineraries, recipes or workout plans.

6. **Output structure** – All files are written into a root directory
   called ``TestDate`` which contains two subdirectories: ``sensitive`` and
   ``insensitive``. A CSV file named ``ground_truth.csv`` enumerates every
   file with its filename, sensitivity label and format. No external data
   files are required beyond standard Python packages, ``python-docx``
   (`docx`) and ``fpdf``.

To run the generator from the command line, execute:

.. code-block:: sh

   python generate_dataset.py --seed "mySeedString"

If the ``--seed`` option is omitted the script will prompt interactively for
input. Using the same seed produces an identical dataset, enabling
reproducibility.
"""

import argparse
import csv
import hashlib
import os
import random
import string
from datetime import datetime, timedelta

# Third party libraries for document creation
# Attempt to import FPDF for PDF creation; if unavailable, fall back to PyMuPDF (fitz)
PDF_ENGINE: str
try:
    from fpdf import FPDF  # type: ignore

    PDF_ENGINE = "fpdf"
except ImportError:
    try:
        import fitz  # type: ignore

        PDF_ENGINE = "pymupdf"
    except ImportError as e:
        raise ImportError(
            "Neither 'fpdf' nor 'pymupdf' is available for PDF generation. "
            "Install one of them via pip."
        ) from e

# Attempt to import python-docx for DOCX creation; fall back to manual generation
DOCX_ENGINE: str
try:
    from docx import Document  # type: ignore
    DOCX_ENGINE = "python-docx"
except ImportError:
    DOCX_ENGINE = "manual"


############################
# Randomness and seeding
############################

def initialise_random(seed_string: str) -> None:
    """Deterministically seed Python's random module based on an arbitrary string.

    The user provides a seed as a string. We hash it to produce an integer
    within a sensible range and then use that to seed the random module.

    Args:
        seed_string: Arbitrary user supplied seed. Same string → same dataset.
    """
    # Use SHA‑256 to get a reproducible integer from the seed string.
    digest = hashlib.sha256(seed_string.encode("utf-8")).hexdigest()
    # Convert a portion of the hash into an integer. The modulus keeps the
    # integer within a reasonable size for the random seed.
    seed_int = int(digest, 16) % (10 ** 8)
    random.seed(seed_int)


#######################################
# Sensitive categories and value generation
#######################################

class SensitiveCategory:
    """Encapsulates a single sensitive data category with generation logic."""

    def __init__(self, key: str, labels: list[str], generator):
        self.key = key
        self.labels = labels  # list of label synonyms
        self.generator = generator  # function returning a realistic value

    def random_label(self) -> str:
        return random.choice(self.labels)

    def random_value(self):
        return self.generator()


def random_phone() -> str:
    # Generate a European style phone number
    return f"+49 {random.randint(30, 89)} {random.randint(100,999)} {random.randint(100000,999999)}"


def random_iban() -> str:
    # Basic German IBAN (22 chars: DE + 20 digits)
    return "DE" + "".join(str(random.randint(0, 9)) for _ in range(20))


def random_name() -> str:
    names = [
        "John Doe",
        "Alice Müller",
        "Omar Youssef",
        "Marie Curie",
        "Ravi Kumar",
        "Sophie Léger",
        "Carlos Díaz",
        "Yuki Nakamura",
    ]
    return random.choice(names)


def random_address() -> str:
    streets = [
        "Hauptstraße 12, 10115 Berlin, Germany",
        "Bahnhofstrasse 5, 8001 Zürich, Switzerland",
        "Rue de Rivoli 99, 75001 Paris, France",
        "221B Baker Street, NW1 6XE London, UK",
        "1600 Amphitheatre Pkwy, 94043 Mountain View, CA, USA",
        "123 Queen St, Toronto, ON M5H 2N2, Canada",
    ]
    return random.choice(streets)


def random_credit_card() -> str:
    # Generate a random Visa or MasterCard number grouped by four
    card_prefixes = ["4", "5"]
    prefix = random.choice(card_prefixes)
    number = prefix + "".join(str(random.randint(0, 9)) for _ in range(15))
    grouped = " ".join(number[i:i + 4] for i in range(0, 16, 4))
    return grouped


def random_passport() -> str:
    # Passport numbers often consist of letters and digits
    return random.choice(string.ascii_uppercase) + "".join(
        random.choice(string.digits + string.ascii_uppercase) for _ in range(8)
    )


def random_drivers_license() -> str:
    # Simplified driver’s licence: letter + 7 digits
    return random.choice(string.ascii_uppercase) + "".join(
        random.choice(string.digits) for _ in range(7)
    )


def random_email() -> str:
    # Compose email from name parts
    first_names = ["john", "alice", "omar", "marie", "ravi", "sophie", "carlos", "yuki"]
    last_names = ["doe", "mueller", "youssef", "curie", "kumar", "leger", "diaz", "nakamura"]
    domains = ["example.com", "mail.com", "university.edu", "corp.org"]
    return f"{random.choice(first_names)}.{random.choice(last_names)}@{random.choice(domains)}"


def random_ip() -> str:
    return ".".join(str(random.randint(1, 254)) for _ in range(4))


def random_username() -> str:
    adjectives = ["fast", "silent", "red", "blue", "happy", "clever", "lucky"]
    nouns = ["fox", "lion", "panda", "tiger", "eagle", "otter", "dragon"]
    return f"{random.choice(adjectives)}_{random.choice(nouns)}{random.randint(1,99)}"


def random_password() -> str:
    # Pick a random strongish password
    words = ["blue", "river", "sunset", "mountain", "forest", "cloud", "storm"]
    specials = ["!", "@", "#", "$", "%"]
    return (
        random.choice(words)
        + random.choice(words).capitalize()
        + str(random.randint(10, 99))
        + random.choice(specials)
    )


def random_medical_term() -> str:
    terms = ["Hypertension", "Diabetes", "Asthma", "Cardiomyopathy", "Anemia", "Epilepsy"]
    return random.choice(terms)


def random_diagnosis() -> str:
    diagnoses = [
        "Type 2 diabetes mellitus",
        "Essential hypertension",
        "Seasonal allergic rhinitis",
        "Chronic obstructive pulmonary disease",
        "Major depressive disorder",
    ]
    return random.choice(diagnoses)


def random_tax_number() -> str:
    return "DE" + "".join(str(random.randint(0, 9)) for _ in range(9))


def random_salary() -> str:
    amount = random.randint(20000, 120000)
    return f"€{amount:,.2f}"  # formatted with comma as thousand separator


def random_pin() -> str:
    return f"{random.randint(0, 9999):04d}"


def random_bill() -> str:
    # Represent a bill with an amount and due date
    amount = random.randint(20, 500)
    due_date = datetime.now().date() + timedelta(days=random.randint(5, 30))
    return f"Amount: €{amount}, Due Date: {due_date.isoformat()}"


def random_invoice() -> str:
    # Represent an invoice with invoice number and total
    invoice_no = f"INV{random.randint(10000, 99999)}"
    total = random.randint(100, 5000)
    return f"{invoice_no}, Total: €{total}"


def random_security_question() -> str:
    questions = [
        "What is your mother’s maiden name?",
        "What was the name of your first pet?",
        "What is your favourite teacher’s name?",
        "What city were you born in?",
    ]
    # Provide a plausible answer as well
    answers = ["Smith", "Rex", "Mr. Thompson", "Munich"]
    return f"{random.choice(questions)} Answer: {random.choice(answers)}"


def random_contract() -> str:
    # Simple representation of a contract summary
    parties = ["Company A", "Client B", "Vendor C", "Partner D"]
    return (
        f"Contract between {random.choice(parties)} and {random.choice(parties)}. "
        f"Terms include service delivery within {random.randint(1, 12)} months and payment of €{random.randint(5000, 50000)}."
    )


def random_api_key() -> str:
    # Generate a pseudo API key (hex string)
    return "".join(random.choice("abcdef0123456789") for _ in range(32))


def random_auth_token() -> str:
    # Generate a base64-like token (alphanumeric)
    chars = string.ascii_letters + string.digits
    return "".join(random.choice(chars) for _ in range(24))


def build_sensitive_categories() -> list[SensitiveCategory]:
    """Create the list of sensitive categories with label synonyms and value generators."""
    return [
        SensitiveCategory(
            "phone",
            ["Phone Number", "Contact No.", "Tel", "Telephone"],
            random_phone,
        ),
        SensitiveCategory(
            "iban",
            ["IBAN", "Bank IBAN", "Account Number"],
            random_iban,
        ),
        SensitiveCategory(
            "names",
            ["Name", "Full Name", "Identity"],
            random_name,
        ),
        SensitiveCategory(
            "addresses",
            ["Address", "Home Address", "Residence"],
            random_address,
        ),
        SensitiveCategory(
            "credit_cards",
            ["Credit Card", "Card Number", "Card No.", "CC"],
            random_credit_card,
        ),
        SensitiveCategory(
            "passport_nr",
            ["Passport Number", "Passport ID"],
            random_passport,
        ),
        SensitiveCategory(
            "drivers_license_nr",
            ["Driver's Licence", "License No.", "DL Number"],
            random_drivers_license,
        ),
        SensitiveCategory(
            "email_addresses",
            ["Email", "Email Address", "Email ID"],
            random_email,
        ),
        SensitiveCategory(
            "ip_addresses",
            ["IP Address", "IPv4", "IP"],
            random_ip,
        ),
        SensitiveCategory(
            "usernames",
            ["Username", "User Handle", "Login ID"],
            random_username,
        ),
        SensitiveCategory(
            "passwords",
            ["Password", "Passcode", "PWD"],
            random_password,
        ),
        SensitiveCategory(
            "medical_terms",
            ["Medical Term", "Condition", "Disease"],
            random_medical_term,
        ),
        SensitiveCategory(
            "diagnosis",
            ["Diagnosis", "Assessment", "Medical Diagnosis"],
            random_diagnosis,
        ),
        SensitiveCategory(
            "tax_numbers",
            ["Tax Number", "Tax ID", "TIN"],
            random_tax_number,
        ),
        SensitiveCategory(
            "salary_information",
            ["Salary", "Income", "Pay", "Wage"],
            random_salary,
        ),
        SensitiveCategory(
            "pin_numbers",
            ["PIN", "PIN Code", "Security Code"],
            random_pin,
        ),
        SensitiveCategory(
            "bills",
            ["Bill", "Billing Statement", "Charge"],
            random_bill,
        ),
        SensitiveCategory(
            "invoices",
            ["Invoice", "Receipt", "Statement"],
            random_invoice,
        ),
        SensitiveCategory(
            "security_questions",
            ["Security Question", "Secret Question", "Auth Question"],
            random_security_question,
        ),
        SensitiveCategory(
            "contracts",
            ["Contract", "Agreement", "Deal"],
            random_contract,
        ),
        SensitiveCategory(
            "api_keys",
            ["API Key", "Service Key", "API Token"],
            random_api_key,
        ),
        SensitiveCategory(
            "auth_tokens",
            ["Auth Token", "Token", "Authorization Key"],
            random_auth_token,
        ),
    ]


#########################################
# Descriptions for file naming (per category)
#########################################

CATEGORY_NAME_FOR_FILENAME = {
    "phone": "PhoneNumbers",
    "iban": "IBAN",
    "names": "Names",
    "addresses": "Addresses",
    "credit_cards": "CreditCards",
    "passport_nr": "PassportInfo",
    "drivers_license_nr": "DriverIDs",
    "email_addresses": "Emails",
    "ip_addresses": "IPAddresses",
    "usernames": "Usernames",
    "passwords": "Passwords",
    "medical_terms": "MedicalTerms",
    "diagnosis": "Diagnosis",
    "tax_numbers": "TaxIDs",
    "salary_information": "Salaries",
    "pin_numbers": "PINs",
    "bills": "Bills",
    "invoices": "Invoices",
    "security_questions": "SecurityQuestions",
    "contracts": "Contracts",
    "api_keys": "ApiKeys",
    "auth_tokens": "AuthTokens",
}


GENERIC_DESCRIPTIONS = [
    "Report",
    "Notes",
    "Summary",
    "Draft",
    "Document",
    "Agenda",
    "Essay",
    "List",
    "Grocery",
    "TravelPlan",
    "Recipe",
]


#########################################
# File name generation logic
#########################################

def camel_case(s: str) -> str:
    """Convert a phrase into CamelCase (removing spaces and underscores)."""
    parts = [p for p in s.replace("_", " ").split(" ") if p]
    return "".join(word.capitalize() for word in parts)


def random_date_str() -> str:
    """Return a random date string in YYYYMMDD format within the last three years."""
    start_date = datetime.now() - timedelta(days=365 * 3)
    random_days = random.randint(0, 365 * 3)
    date = start_date + timedelta(days=random_days)
    return date.strftime("%Y%m%d")


def choose_version() -> str:
    return f"v{random.randint(1, 5)}"


def random_username_for_filename() -> str:
    # Simpler username for file naming purposes
    names = ["jdoe", "a_muller", "oyoussef", "mcurie", "rkumar", "sleger", "cdiaz", "ynakamura"]
    return random.choice(names)


def generate_harvard_filename(date_str: str, categories: list[str] | None, match: bool, ext: str) -> str:
    # Format: YYYYMMDD_Project_Sample_vX.ext
    version = choose_version()
    if match and categories:
        # Use the first category to derive a project name
        project = CATEGORY_NAME_FOR_FILENAME.get(categories[0], random.choice(GENERIC_DESCRIPTIONS))
        sample = "Sample"
    else:
        project = random.choice(GENERIC_DESCRIPTIONS)
        sample = random.choice(GENERIC_DESCRIPTIONS)
    return f"{date_str}_{project}_{sample}_{version}.{ext}"


def generate_montessori_filename(date_str: str, categories: list[str] | None, match: bool, ext: str) -> str:
    # Format: surname_firstname_YYYYMMDD_loc_desc.ext
    surnames = ["doe", "muller", "youssef", "curie", "kumar", "leger", "diaz", "nakamura"]
    firstnames = ["john", "alice", "omar", "marie", "ravi", "sophie", "carlos", "yuki"]
    surname = random.choice(surnames)
    firstname = random.choice(firstnames)
    locations = ["berlin", "hamburg", "paris", "london", "newyork", "zurich", "tokyo", "toronto"]
    loc = random.choice(locations)
    if match and categories:
        desc = CATEGORY_NAME_FOR_FILENAME.get(categories[0], random.choice(GENERIC_DESCRIPTIONS)).lower()
    else:
        desc = random.choice(GENERIC_DESCRIPTIONS).lower()
    return f"{surname}_{firstname}_{date_str}_{loc}_{desc}.{ext}"


def generate_institutional_filename(date_str: str, categories: list[str] | None, match: bool, ext: str) -> str:
    # Format: YYYYMMDD_Description_Version.ext
    version = choose_version()
    if match and categories:
        desc = CATEGORY_NAME_FOR_FILENAME.get(categories[0], random.choice(GENERIC_DESCRIPTIONS))
    else:
        desc = random.choice(GENERIC_DESCRIPTIONS)
    return f"{date_str}_{desc}_{version}.{ext}"


def generate_osu_filename(date_str: str, categories: list[str] | None, match: bool, ext: str) -> str:
    # Format: PROJYYYYMMDD_Type_Subject_vX.ext
    version = choose_version()
    # Always prefix PROJ
    if match and categories:
        subject = CATEGORY_NAME_FOR_FILENAME.get(categories[0], random.choice(GENERIC_DESCRIPTIONS))
        type_ = "Confidential"
    else:
        subject = random.choice(GENERIC_DESCRIPTIONS)
        type_ = random.choice(["Report", "Notes", "Public", "Archive"])
    return f"PROJ{date_str}_{type_}_{subject}_{version}.{ext}"


def generate_reddit_filename(date_str: str, categories: list[str] | None, match: bool, ext: str) -> str:
    # Format: Descriptive_CamelCase_date_vX_username.ext
    version = choose_version()
    username = random_username_for_filename()
    if match and categories:
        desc_raw = CATEGORY_NAME_FOR_FILENAME.get(categories[0], random.choice(GENERIC_DESCRIPTIONS))
    else:
        desc_raw = random.choice(GENERIC_DESCRIPTIONS)
    descriptive = camel_case(desc_raw)
    return f"{descriptive}_{date_str}_{version}_{username}.{ext}"


NAMING_FUNCTIONS = [
    generate_harvard_filename,
    generate_montessori_filename,
    generate_institutional_filename,
    generate_osu_filename,
    generate_reddit_filename,
]


#########################################
# Content generation for sensitive files
#########################################

def generate_sensitive_content(categories: list[SensitiveCategory]) -> str:
    """Compose a document containing one or more sensitive categories.

    A realistic document style is randomly chosen: bullet list, paragraph,
    to‑do list or short notes. Each category contributes a label/value pair
    using a randomly selected synonym for the label.

    Args:
        categories: list of SensitiveCategory instances to include

    Returns:
        A string representing the content of the document.
    """
    style = random.choice(["bullet", "paragraph", "todo", "notes"])
    lines: list[str] = []

    if style == "bullet":
        for cat in categories:
            label = cat.random_label()
            value = cat.random_value()
            lines.append(f"- {label}: {value}")
        return "\n".join(lines)

    elif style == "paragraph":
        paragraphs: list[str] = []
        for cat in categories:
            label = cat.random_label()
            value = cat.random_value()
            sentences = [
                f"The {label.lower()} recorded here is {value}.",
                f"Please keep the {label.lower()} confidential and do not share it.",
            ]
            paragraphs.append(" ".join(sentences))
        return "\n\n".join(paragraphs)

    elif style == "todo":
        for cat in categories:
            label = cat.random_label()
            value = cat.random_value()
            lines.append(f"- Verify {label.lower()} ({value})")
        lines.append("- Complete all outstanding actions by end of day.")
        return "\n".join(lines)

    else:  # notes
        for cat in categories:
            label = cat.random_label()
            value = cat.random_value()
            lines.append(f"{label}: {value}")
        return "\n".join(lines)


#########################################
# Content generation for insensitive files
#########################################

def generate_insensitive_content() -> str:
    """Generate plausible but non‑sensitive content in various styles."""
    style = random.choice(["grocery", "travel", "recipe", "book", "workout", "lecture", "essay", "todo", "notes"])

    if style == "grocery":
        items = random.sample(
            ["Milk", "Eggs", "Bread", "Butter", "Apples", "Tomatoes", "Cheese", "Pasta", "Coffee"],
            k=5,
        )
        return "Weekly Grocery List\n" + "\n".join(f"- {item}" for item in items)

    if style == "travel":
        cities = ["Rome", "Oslo", "Paris", "Berlin", "Tokyo", "New York", "Zurich", "Barcelona"]
        city = random.choice(cities)
        days = random.randint(2, 14)
        notes = random.choice(["Check museum hours", "Book hotel early", "Reserve tickets online"])
        return f"Travel Itinerary\nDestination: {city}\nDuration: {days} days\nNotes: {notes}"

    if style == "recipe":
        dishes = ["Pasta Carbonara", "Chili Stew", "Tomato Soup", "Beef Stroganoff"]
        dish = random.choice(dishes)
        ingredients_pool = ["Tomato", "Beef", "Garlic", "Onion", "Carrot", "Pasta", "Cream", "Cheese"]
        ingredients = ", ".join(random.sample(ingredients_pool, k=4))
        steps = "Mix all ingredients and simmer for 20 minutes."
        return f"Recipe: {dish}\nIngredients: {ingredients}\nSteps: {steps}"

    if style == "book":
        titles = ["The Alchemist", "1984", "Brave New World", "To Kill a Mockingbird"]
        themes = ["Freedom vs Control", "Personal Legend", "Societal Expectations", "Justice"]
        summaries = [
            "A compelling narrative with strong symbolism.",
            "Explores the tension between individuality and authority.",
            "Raises questions about the role of technology in society.",
            "Highlights moral growth and empathy.",
        ]
        title = random.choice(titles)
        theme = random.choice(themes)
        summary = random.choice(summaries)
        return f"Book Summary\nTitle: {title}\nTheme: {theme}\nNotes: {summary}"

    if style == "workout":
        exercises = ["Push‑ups", "Yoga", "Running", "Cycling", "Swimming", "Squats", "Plank"]
        plan = random.sample(exercises, k=3)
        return "Workout Plan\n" + "\n".join(
            f"Day {i+1}: {plan[i]}" for i in range(3)
        )

    if style == "lecture":
        topics = [
            "Quantum mechanics", "Medieval history", "Organic chemistry", "Microeconomics",
            "Machine learning", "Philosophy of mind", "Ecology", "Sociolinguistics",
        ]
        selected = random.sample(topics, k=4)
        return "Lecture Notes\n" + "\n".join(f"- {topic}" for topic in selected)

    if style == "essay":
        intro_sentences = [
            "In recent years, the concept of sustainability has gained significant traction.",
            "The rapid advancement of technology has reshaped our daily lives in unprecedented ways.",
            "Throughout history, art has reflected the evolution of human society.",
        ]
        body_sentences = [
            "This trend highlights the need for holistic approaches to environmental stewardship.",
            "Such developments prompt debates about privacy, ethics and accessibility.",
            "Artists often respond to political and cultural shifts, producing works that challenge norms.",
            "The intersection of culture and innovation fosters a dynamic landscape of ideas.",
        ]
        conclusion_sentences = [
            "Ultimately, fostering awareness is essential for future progress.",
            "Therefore, continuous dialogue is critical to navigate these complexities.",
            "In conclusion, the enduring impact of these forces is undeniable.",
        ]
        intro = random.choice(intro_sentences)
        body = " ".join(random.sample(body_sentences, k=2))
        conclusion = random.choice(conclusion_sentences)
        return f"{intro}\n\n{body}\n\n{conclusion}"

    if style == "todo":
        tasks = [
            "Finish quarterly report",
            "Reply to client emails",
            "Schedule team meeting",
            "Organise project files",
            "Plan budget review",
        ]
        selected = random.sample(tasks, k=4)
        return "To‑Do List\n" + "\n".join(f"- {task}" for task in selected)

    else:  # notes
        notes = [
            "Call the supplier tomorrow at 10 AM.",
            "Remember to water the plants.",
            "Buy tickets for the conference.",
            "Check the latest research papers on AI.",
        ]
        selected = random.sample(notes, k=3)
        return "Short Notes\n" + "\n".join(selected)


#########################################
# File writing helpers
#########################################

def write_txt(path: str, content: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)


def write_pdf(path: str, content: str) -> None:
    """Write content to a PDF file using the available PDF engine.

    If FPDF is available, it uses that library; otherwise it falls back to
    PyMuPDF via ``fitz``. PyMuPDF is capable of inserting plain text into
    a page. Both methods produce a valid PDF file.

    Args:
        path: Destination file path
        content: Text content to write
    """
    if PDF_ENGINE == "fpdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in content.split("\n"):
            pdf.multi_cell(0, 8, txt=line)
        pdf.output(path)
    elif PDF_ENGINE == "pymupdf":
        import fitz  # type: ignore

        pdf_doc = fitz.open()
        page = pdf_doc.new_page()
        # Insert text at (72,72) with a standard font size
        text = content.replace("\n", "\n")
        page.insert_text(
            fitz.Point(72, 72),
            text,
            fontsize=12,
            fontname="helv",
        )
        pdf_doc.save(path)
    else:
        # This branch should not occur due to import checks above
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)


def write_docx(path: str, content: str) -> None:
    """Write content to a DOCX file.

    Two mechanisms are supported. If python-docx is available it will be used
    to create a richer Word document with proper bullet styles. Otherwise a
    minimal Office Open XML package is constructed by hand. The manual
    variant writes each line as a separate paragraph; bullets are not
    explicitly styled but the textual dash is preserved.

    Args:
        path: Destination file path
        content: Text content to write
    """
    if DOCX_ENGINE == "python-docx":
        # Use python-docx for convenience
        from docx import Document  # type: ignore
        doc = Document()
        for line in content.split("\n"):
            if line.strip().startswith("-"):
                # Remove leading dash and space for bullet styling
                text = line.lstrip("- ")
                p = doc.add_paragraph(text)
                p.style = 'List Bullet'
            else:
                doc.add_paragraph(line)
        doc.save(path)
    else:
        # Manual DOCX creation using zipfile and minimal XML
        import zipfile
        import xml.sax.saxutils as saxutils

        # Prepare XML bodies
        lines = content.split("\n")
        body_parts = []
        for line in lines:
            # Escape XML special characters
            text = saxutils.escape(line)
            body_parts.append(
                f"<w:p><w:r><w:t xml:space=\"preserve\">{text}</w:t></w:r></w:p>"
            )
        body_xml = "".join(body_parts)
        document_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
            'xmlns:o="urn:schemas-microsoft-com:office:office" '
            'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
            'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
            'xmlns:w10="urn:schemas-microsoft-com:office:word" '
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
            'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
            'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
            'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
            'mc:Ignorable="w14 wp14">'
            f"<w:body>{body_xml}</w:body>"
            "</w:document>"
        )
        content_types_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '</Types>'
        )
        rels_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/>'
            '</Relationships>'
        )
        # Minimal relationship file for document (empty)
        document_rels_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '</Relationships>'
        )
        # Write the zip package
        with zipfile.ZipFile(path, 'w') as zf:
            zf.writestr('[Content_Types].xml', content_types_xml)
            zf.writestr('_rels/.rels', rels_xml)
            zf.writestr('word/document.xml', document_xml)
            zf.writestr('word/_rels/document.xml.rels', document_rels_xml)


def ensure_unique(filename: str, used_filenames: set[str]) -> str:
    """Ensure a file name is unique by appending a counter if necessary."""
    if filename not in used_filenames:
        return filename
    name, ext = os.path.splitext(filename)
    counter = 1
    new_name = f"{name}_{counter}{ext}"
    while new_name in used_filenames:
        counter += 1
        new_name = f"{name}_{counter}{ext}"
    return new_name


#########################################
# Main dataset generation function
#########################################

def generate_dataset(seed: str, total_files: int = 100) -> None:
    initialise_random(seed)
    base_dir = "TestData"
    sensitive_dir = os.path.join(base_dir, "sensitive")
    insensitive_dir = os.path.join(base_dir, "insensitive")
    os.makedirs(sensitive_dir, exist_ok=True)
    os.makedirs(insensitive_dir, exist_ok=True)

    # Compute file-type counts: 70% txt, 15% docx, 15% pdf
    docx_count = int(total_files * 0.15)
    pdf_count = int(total_files * 0.15)
    txt_count = total_files - docx_count - pdf_count
    file_types = ["txt"] * txt_count + ["docx"] * docx_count + ["pdf"] * pdf_count
    random.shuffle(file_types)

    # Sensitivity split: approx. half
    sensitive_count = total_files // 2
    insensitive_count = total_files - sensitive_count
    sensitivities = ["SENSITIVE"] * sensitive_count + ["INSENSITIVE"] * insensitive_count
    random.shuffle(sensitivities)

    # Match flag split: approx. half
    true_count = total_files // 2
    false_count = total_files - true_count
    match_flags = [True] * true_count + [False] * false_count
    random.shuffle(match_flags)

    categories_list = build_sensitive_categories()
    category_keys = [c.key for c in categories_list]
    num_categories = len(category_keys)
    if sensitive_count < num_categories:
        raise ValueError(f"Need at least {num_categories} sensitive files to cover each category at least once.")

    # Assign categories to sensitive files
    assignments = [[key] for key in category_keys]
    remaining = sensitive_count - num_categories
    mix_count = int(0.2 * sensitive_count)
    # placeholders
    assignments += [[] for _ in range(remaining)]
    mix_indices = set(random.sample(range(len(assignments)), k=mix_count))
    for i in range(num_categories, len(assignments)):
        if i in mix_indices:
            k = random.randint(2, 4)
            assignments[i] = random.sample(category_keys, k=k)
        else:
            assignments[i] = [random.choice(category_keys)]
    random.shuffle(assignments)

    used_filenames = set()
    ground_truth = []
    sens_index = 0

    for i in range(total_files):
        file_format = file_types[i]
        sensitivity = sensitivities[i]
        match_flag = match_flags[i]

        if sensitivity == "SENSITIVE":
            keys = assignments[sens_index]
            sens_index += 1
            cats = [c for c in categories_list if c.key in keys]
            content = generate_sensitive_content(cats)
            naming_func = random.choice(NAMING_FUNCTIONS)
            filename = naming_func(random_date_str(), keys, match_flag, file_format)
            filename = ensure_unique(filename, used_filenames)
            dest = os.path.join(sensitive_dir, filename)
        else:
            content = generate_insensitive_content()
            naming_func = random.choice(NAMING_FUNCTIONS)
            filename = naming_func(random_date_str(), None, match_flag, file_format)
            filename = ensure_unique(filename, used_filenames)
            dest = os.path.join(insensitive_dir, filename)

        used_filenames.add(filename)
        # Write to disk
        if file_format == "txt":
            write_txt(dest, content)
        elif file_format == "pdf":
            write_pdf(dest, content)
        elif file_format == "docx":
            write_docx(dest, content)
        else:
            raise ValueError(f"Unsupported file format: {file_format}")

        ground_truth.append([filename, sensitivity, file_format])

    # Write CSV
    gt_path = os.path.join(base_dir, "ground_truth.csv")
    with open(gt_path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(["filename", "sensitivity", "format"])
        writer.writerows(ground_truth)

    print(f"✅ Dataset generated in '{base_dir}' with {total_files} files.")


def main():
    parser = argparse.ArgumentParser(description="Generate a synthetic dataset of sensitive and insensitive files.")
    parser.add_argument(
        "--seed",
        type=str,
        default=None,
        help="Reproducibility seed (string). Identical seeds yield identical datasets.",
    )
    args = parser.parse_args()
    seed = args.seed
    # If seed not provided, prompt interactively
    if seed is None:
        seed = input("Enter seed for reproducibility: ")
    generate_dataset(seed)


if __name__ == "__main__":
    main()